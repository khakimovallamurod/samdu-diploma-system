import pandas as pd
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from datetime import date 


def get_student_data(file_path):
    df = pd.read_excel(file_path, header=7)

    selected_columns = df.iloc[:, [0, 1, 9, 12, 13, 14]]
    selected_columns.columns = ["ID", 'F.I.Sh.', 'Yonalish', "diplom_id", "tartib_raqam", "Kvalifikatsiya"]

    selected_columns = selected_columns.dropna(
        how='all',
        subset=["ID", 'F.I.Sh.', 'Yonalish', 'Kvalifikatsiya', 
                "diplom_id", "tartib_raqam"]
    )

    selected_columns = selected_columns.fillna("")

    return selected_columns.to_dict(orient='records')

def date_to_string(date_obj):
    day = date_obj.day
    month = date_obj.month
    year = date_obj.year
    month_names = [
        "Yanvar", "Fevral", "Mart", "Aprel", "May", "Iyun",
        "Iyul", "Avgust", "Sentabr", "Oktabr", "Noyabr", "Dekabr"
    ]
    month_name = month_names[month - 1]
    return f"{year} yil {day}-{month_name.lower()}dagi"


def create_diplom_kuchirma_hujjat(student_data, sana,  fayl_nomi='bitiruvchi_diplom_kuchirma.docx'):
    doc = Document()

    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(8.50)   # A4 gorizontal eni
    section.page_height = Inches(11.0)  # A4 gorizontal bo‘yi
    section.top_margin = Inches(0.59)
    section.bottom_margin = Inches(0.39)
    section.left_margin = Inches(0.59)
    section.right_margin = Inches(0.59)

    def add_line(cell, text, size=12, bold=False, center=True, left_indent_cm=0):
        p = cell.add_paragraph()
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        run._element.rPr.rFonts.set(qn('w:ascii'), 'Times New Roman')
        run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Times New Roman')
        run.font.size = Pt(size)
        run.bold = bold
        p_format = p.paragraph_format
        p_format.line_spacing = Pt(12)
        p_format.space_before = Pt(4)
        # p_format.space_after = Pt(4)

        if center:
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        else:
            p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            p_format.left_indent = Cm(left_indent_cm)

        

    def add_line_mixed(cell, parts, center=True, left_indent_cm=0):
        p = cell.add_paragraph()
        p_format = p.paragraph_format
        p_format.line_spacing = Pt(12)
        # p_format.space_before = Pt(4)
        # p_format.space_after = Pt(4)
        if center:
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        else:
            p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            p_format.left_indent = Cm(left_indent_cm)


        for text, bold, size in parts:
            run = p.add_run(text)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(size)
            run.bold = bold
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
            run._element.rPr.rFonts.set(qn('w:ascii'), 'Times New Roman')
            run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Times New Roman')

    def add_empty_line(cell, size=6):
        p = cell.add_paragraph()
        run = p.add_run(" ")  # faqat bitta bo‘sh joy, matnsiz
        run.font.name = 'Times New Roman'
        run.font.size = Pt(size)
        
        p.paragraph_format.line_spacing = Pt(12)
        # p.space_before = Pt(2)  
        # p.space_after = Pt(2)
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    def fill_cell(cell, item):
        for para in cell.paragraphs:
            p = para._element
            p.getparent().remove(p)
        add_line(cell, "O‘ZBEKISTON RESPUBLIKASI", 20, bold=True)
        add_line(cell, "BAKALAVR", 20, bold=True)
        add_line(cell, "DIPLOMI", 20, bold=True)
        add_empty_line(cell, 9)
        qaror_id = item['diplom_id'] if item['diplom_id'] != '' else "B №____________"
        add_line(cell, f'  {qaror_id}', 16)
        add_empty_line(cell, 9)
        add_line(cell, "K O‘ Ch I R M A", 26, bold=True)
        add_line(cell, "SAMARQAND DAVLAT UNIVERSITETI", 18, bold=True)
        add_empty_line(cell, 14)
        add_line(cell, "Davlat attestasiya komissiyasining", 14, bold=True)
        add_line(cell, f"{sana}", 14, bold=True)
        add_line(cell, "qaroriga binoan", 14, bold=True)
        add_empty_line(cell, 14)
        fish = item['F.I.Sh.'] if item['F.I.Sh.'] != '' else "____________________________________"
        add_line_mixed(cell, [
            (fish.upper(), True, 16),
            ("ga", False, 14)
        ])
        add_empty_line(cell, 11)
        yunalish = item['Yonalish'] if item['Yonalish'] != '' else "____________________________________"
        add_line(cell, f"{yunalish}", 18)
        add_line(cell, f"  yo’nalishi bo’yicha", 14)
        add_empty_line(cell, 11)
        add_line(cell, "B A K A L A V R", 26, bold=True)
        add_line(cell, "DARAJASI", 22, bold=True)
        add_empty_line(cell, 8)
        kval = item['Kvalifikatsiya'] if item['Kvalifikatsiya'] != '' else "____________________________________"
        add_line_mixed(cell, [
            ('va ', False, 18),
            (kval, True, 18)
        ])
        add_line(cell, f"kvalifikatsiyasi berildi", 18)
        add_empty_line(cell, 10)
        tartib_raqam = item['tartib_raqam'] if item['tartib_raqam'] != '' else "________________"
        add_line_mixed(cell, [
            ('Ro‘yxatga olish raqami: ', False, 16),
            (f'{tartib_raqam}', True, 16)
        ], center=False, left_indent_cm=1.1)
        add_empty_line(cell, 9)
        id = item['ID'] if item['ID'] != '' else "________"
        add_line(cell, f"Ushbu ko‘chirma faqat {id}-sonli yo‘llanma bilan o‘z kuchiga ega.", 16, center=False, left_indent_cm=1.1)
        add_empty_line(cell, 9)
        add_line(cell, "Davlat attestatsiya va taqsimot", 16, center=False, left_indent_cm=1.1)
        add_line(cell, "komissiyalari raisi.", 16, center=False, left_indent_cm=1.1)
        add_empty_line(cell, 9)
        add_line(cell, "Rektor      						R.I.Xalmuradov", 16, bold=True, left_indent_cm=3)
        
    # Talabalarni juftlab sahifalarga joylashtirish
    for i in range(len(student_data)):
        table = doc.add_table(rows=1, cols=1)
        row = table.rows[0]
        cell = row.cells[0]

        fill_cell(cell, student_data[i])
        doc.add_page_break()

    doc.save(fayl_nomi)
    return fayl_nomi




def main(file_path, sana):
    sana_str = date_to_string(sana)
    student_data = get_student_data(file_path)
    output_file = create_diplom_kuchirma_hujjat(student_data, sana_str)
    return output_file

if __name__ == "__main__":
    file_path = '2025 yil 10 iyun.xlsx'  # Replace with your actual file path
    sana = date(2023, 10, 1)  
    main(file_path, sana)
